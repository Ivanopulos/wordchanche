#pyinstaller --onefile xltoword.py cd C:\Users\IMatveev\PycharmProjects\wordchanche\
# правильный наташин
# требует
## все эксели закрыты
## выбираемого шаблонв ворд
## наличие шаблон написания справки.хлсх и в нем вляющей ячейки на вкладке Таб  2019 B1
## установленной надстройки ивах с доступным макросом цвета (будет обращаться к документу шаблон пмо)
## наличие/нет столбцов \\if 'естьнадстройка' \\if 'color' \\'regions'
#\н дежурный макрос хлсм
#\н пересмотреть привязку на шаблон пмо в сторону более универсального
import os
import zipfile
import re
import numpy as np
import pandas  # +openpyxl
from tkinter import filedialog
import sys
import win32com.client  # pip install pypiwin32
import docx #pip install python-docx
import shutil
import datetime
import time
from datetime import datetime
def dtm():  # функция эхо времени(надо ее вызвать в начале программы и в любое время в дальнейшем для вывода времени и времени выполнения программы)
    if not hasattr(dtm, 'dt0'):
        dtm.dt0 = datetime.now()
    else:
        dt1 = datetime.now()
        print(dt1-dtm.dt0, "//", dt1)
dtm()
print(os.environ.get( "USERNAME" ))
def run_macro(name, new_value, df):
    print('macro', name)
    if os.path.exists(name):
        xl = win32com.client.Dispatch("Excel.Application")
        #xl.DisplayAlerts = True
        #xl.Visible = True
        wb = xl.Workbooks.Open(Filename=name, ReadOnly=0)
        if 'естьнадстройка' in df.columns:
            xl.Application.Run("'C:\\Users\\" + os.environ.get("USERNAME") + "\\AppData\\Roaming\\Microsoft\\AddIns\\Ivax.xlam'!UpdateFormulas", wb)
        #time.sleep(20)
        #for sheet in wb.Sheets:
        #    print(sheet.Name)
        if not new_value == 'ъъ':#под столбец регионс
            ws = wb.Worksheets("Таб  2019")
            ws.Range("B1").Value = new_value
        else:
            ws = wb.Worksheets(1)
            #print(ws.Range("B6").Value)#, xl.AddIns("Ivax").Installed)
        if 'color' in df.columns:
            xl.Application.Run("'C:\\Users\\" + os.environ.get( "USERNAME" ) + "\\AppData\\Roaming\\Microsoft\\AddIns\\Ivax.xlam'!цвета")
        wb.Close(SaveChanges=True)
        xl.Application.Quit()
        time.sleep(1)
        del xl
    # print('File refreshed!')
def resource_path(relative):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative)
    return os.path.join(relative)
def delete_paragraph(paragraph):   # Delete stroke+
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
print("ворд")
pathword = filedialog.askopenfilename()
pathwork = os.path.dirname(pathword)
pathzip = pathwork + "/B.zip"
pathword2 = pathwork + "/Шаблон написания справки.xlsx"
#/n сделать хлсм с запуском пускового макроса
mem = 0
memcol = 0
isch = 0
usch = 0
found = ""

try:  # поиск списка регионов, если нет - ъъ идет как отсутствме замены
    regions = pandas.read_excel(pathword2).dropna(subset=['regions'])
except:
    regions = pandas.DataFrame()
if 'regions' not in regions.columns:#40305 40528
    regions['regions'] = pandas.NA # Создаем столбец 'regions' с пустыми значениями
    regions.loc[0, 'regions'] = 'ъъ'

for region in regions['regions']:
    mem = 0
    memcol = 0
    isch = 0
    usch = 0
    found = ""
    df = pandas.read_excel(pathword2, dtype=str)
    print(df.loc[1:17, ['metka', 'chenge']])
    # if 'нецвет' in df.columns:
    #     pass
    # else:
    run_macro(pathword2, region, df)  # по душу макроса цвета и перебора регионов (нафига вот в одной куче)

    doc = docx.Document(pathword)
    # for para in doc.paragraphs:
    #     delete = False
    #     for run in para.runs:
    #         if len(run.text) > 0:
    #             if run.text[0] == "{":
    #                 nasr = df[df["metka"] == run.text]["chenge"].to_string(header=False, index=False)
    #                 if nasr != "NaN":
    #                     delete = False
    #                     break
    #     if delete:#and para.text[0]=='{' and para.text[-1]=='}'
    #         print(para.text, "Deleted")
    #         delete_paragraph(para)
    for para in doc.paragraphs:
        markers_empty = True
        found_marker = False
        for run in para.runs:
            pos = 0
            while "{" in run.text[pos:]:  # Проверяем все вхождения "{"
                found_marker = True
                start = run.text.find('{', pos)
                end = run.text.find('}', start) + 1
                if end > start:
                    marker_text = run.text[start:end]
                    #print("\\"+run.text[start:end]+"\\")
                    nasr = df[df["metka"] == marker_text]["chenge"].to_string(header=False, index=False)
                    print(nasr)
                    if nasr != "NaN":
                        markers_empty = False
                        break  # Прерываем цикл, если найдена непустая метка
                pos = end  # Продолжаем поиск со следующей позиции
            if not markers_empty:
                break  # Прерываем цикл, если в одном из run найдена непустая метка

        if markers_empty and found_marker:  # Удаляем параграф, если все метки пусты и хотя бы одна метка была найдена
            print(para.text, "Deleted")
            delete_paragraph(para)
    doc.save(pathwork + "/B.docx")
    try:
        os.remove(pathwork + "/B.zip")
    except:
        asd = 1
    df = df.fillna("")
    os.rename(pathwork + "/B.docx", pathwork + "/B.zip")
    fantasy_zip = zipfile.ZipFile(pathzip)  # extract zip (+need rename docx to zip +need raname vise versa
    fantasy_zip.extractall(pathwork + "/B")
    fantasy_zip.close()
    with open(pathwork + "/B/word/document.xml", 'r', encoding='utf-8') as f:  # save before chenge
        get_all = f.readlines()
    print("xml opened")
    with open(pathwork + "/B/word/document.xml", 'w', encoding='utf-8') as f:  # look for { and chenge it
        for i in get_all:         # STARTS THE NUMBERING FROM 1 (by default it begins with 0)
            usch = len(i)-1
            for u in i:
                try:
                    if get_all[isch][usch] == "}":
                        mem = 1
                        memcol = 0
                except:
                    print(isch, usch, u, i, get_all)
                    print(get_all[isch][usch])
                if memcol == 1: #замена цвета
                    if get_all[isch][usch:usch+7] == "w:fill=":
                        memcol = 0
                        if not get_all[isch][usch+8] == "a":
                            dl = 6
                            #print(get_all[isch][usch - 2:usch + 40])
                            get_all[isch] = get_all[isch][:usch+8] + str(col) + get_all[isch][usch + 8 + dl:]
                        else:
                            dl = 4
                            get_all[isch] = get_all[isch][:usch + 8] + str(col) + get_all[isch][usch + 8 + dl:]
                            #print(get_all[isch][usch-2:usch+40])
                        # print(get_all[isch][usch:usch+50])
                if mem == 1:
                    found = get_all[isch][usch] + found
                if get_all[isch][usch] == "{":
                    mem = 0
                    #print(found, 98)
                    #found=re.sub(r'(\{\d+)<[^}{]+>(\d+\})', r'\1\2', found) лимитировать б длину 5 сиволами
                    print(found, 99)

                    tx = df[df["metka"] == found]["chenge"].values[0]#header=False,

                    try:
                        float(tx)
                        tx = str(tx).replace(".", ",")
                    except:
                        asd = 0


                    if 'color' in df.columns:
                        col = df[df["metka"] == found]["color"].values[0]
                        if col == "NaN" or col == "FFFFFF" or col == "":
                            pass
                        else:
                            memcol = 1
                    get_all[isch] = get_all[isch][:usch] + tx + get_all[isch][usch + len(found):]
                    found = ""
                usch = usch - 1
            isch = isch + 1
        f.writelines(get_all)
    print("XML chanched")
    try:
        os.remove(pathwork + "/B.zip")
    except:
        asd = 1
    fantasy_zip = zipfile.ZipFile(pathwork + "/B.zip", 'w')
    for folder, subfolders, files in os.walk(pathwork + "/B"):
        for file in files:
            fantasy_zip.write(os.path.join(folder, file), os.path.relpath(os.path.join(folder, file), pathwork + "/B"))
    fantasy_zip.close()  # transform it to zip
    print("zip saved")
    name = str(df.iloc[0, 1])
    if len(name) > 82:
        name = "О разв ПМСП в " + name[82:]
        print(name)
    else:
        name = "Документ " + str(datetime.today().date())
    try:
        os.remove(pathwork + "/" + name + ".docx")
        print(name, "removed")
    except:
        asd = 1

    # os.rename(pathwork + "/B.zip", pathwork + "/" + datetime.now().strftime("%d.%m.%y") + pathword.split("/")[-1].split(".")[0] + "/" + name + ".docx")
    # Формируем путь к новому каталогу
    new_dir = os.path.join(pathwork, datetime.now().strftime("%d.%m.%y") + pathword.split("/")[-1].split(".")[0])

    # Создаем каталог, если он еще не существует
    os.makedirs(new_dir, exist_ok=True)

    # Строим новый путь к файлу после переименования
    new_path = os.path.join(new_dir, name + ".docx")

    # Переименовываем (перемещаем) файл
    # роняет если он есть os.rename(os.path.join(pathwork, "B.zip"), new_path)
    shutil.move(os.path.join(pathwork, "B.zip"), new_path)

    shutil.rmtree(pathwork + "/B/")
    print("FINISH", region)
    dtm()
